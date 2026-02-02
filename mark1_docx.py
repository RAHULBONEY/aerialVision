# pip install python-docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('AERIAL VISION - MARK 1 MODEL REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Basic Info
doc.add_heading('Model Overview', level=1)
doc.add_paragraph(f"File: best_v1_daytime.pt")
doc.add_paragraph(f"Date: January 29, 2026")
doc.add_paragraph(f"Size: 49.61 MB")

# Architecture
doc.add_heading('Architecture', level=1)
specs = [
    ("Base Model", "YOLOv8 Medium (YOLOv8m)"),
    ("Parameters", "25,860,373 (25.9M)"),
    ("Layers", "169"),
    ("FLOPs", "79.1 GFLOPs"),
    ("Input Size", "640×640 px"),
    ("Task", "Object Detection")
]
for key, value in specs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"{key}: ").bold = True
    p.add_run(value)

# Classes
doc.add_heading('Detection Classes (7 Total)', level=1)
classes = ['sedan', 'minibus', 'truck', 'pickup', 'bus', 'cement truck', 'trailer']
for i, cls in enumerate(classes):
    doc.add_paragraph(f"{i}. {cls.title()}", style='List Number')

# Dataset
doc.add_heading('Training Data', level=1)
doc.add_paragraph("Total Images: 1,948")
doc.add_paragraph("Training: 1,557 images")
doc.add_paragraph("Validation: 391 images")

# Performance
doc.add_heading('Performance Metrics', level=1)
metrics = [
    ("Precision", "95.02%"),
    ("Recall", "94.23%"),
    ("mAP@50", "97.80%"),
    ("mAP@50-95", "79.41%")
]
for metric, value in metrics:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f"{metric}: ").bold = True
    p.add_run(value)

# Save
doc.save('Mark1_Model_Report.docx')
print("✅ Report saved as 'Mark1_Model_Report.docx'")